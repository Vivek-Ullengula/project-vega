"""Universal local chunking pipeline for Project Vega knowledge bases.

The pipeline is intentionally offline: it parses local files or websites, cleans
the extracted text, adds stable citation metadata, and writes one Markdown file
per chunk. Upload the output folder to S3 later and create/sync the Bedrock KB
data source with chunking set to NONE.

Design rules:
- Do not mutate the current production chunks.
- Preserve source wording; LLM/Agno chunkers may choose boundaries only.
- Keep enough locator metadata for future source highlighting:
  page numbers, sheet/row ranges, source URL, file path, block index, char span.
- Prefer deterministic structure-aware chunking for manuals, websites, and
  tables. Use Agno/Docling as optional parser/chunker when installed.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import html
import json
import os
import re
import shutil
import stat
import zipfile
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from datetime import UTC, datetime
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Iterable
from urllib.parse import urlparse
from urllib.request import Request, urlopen
import xml.etree.ElementTree as ET


DEFAULT_SOURCE_PATHS = (
    Path("data/bedrock_ingest"),
    Path("data/guide_sections"),
    Path("data/full_manuals"),
)
DEFAULT_OUTPUT_DIR = Path("scratch/kb_chunks_universal")
DEFAULT_MAX_CHARS = 3600
MIN_BODY_CHARS = 24

PROPERTY_MANUAL_URL = "https://bindingauthority.coactionspecialty.com/manuals/property.html"
GL_GUIDE_MANUAL_URL = "https://bindingauthority.coactionspecialty.com/manuals/guide.html"

TEXT_SUFFIXES = {".txt", ".text", ".rst"}
MARKDOWN_SUFFIXES = {".md", ".markdown"}
HTML_SUFFIXES = {".html", ".htm"}
CSV_SUFFIXES = {".csv", ".tsv"}
EXCEL_SUFFIXES = {".xlsx", ".xlsm", ".xltx", ".xltm"}
DOCX_SUFFIXES = {".docx"}
PDF_SUFFIXES = {".pdf"}
JSON_SUFFIXES = {".json"}
CODE_SUFFIXES = {
    ".py",
    ".js",
    ".jsx",
    ".ts",
    ".tsx",
    ".java",
    ".go",
    ".rs",
    ".cs",
    ".cpp",
    ".c",
    ".h",
    ".sql",
}
SUPPORTED_SUFFIXES = (
    TEXT_SUFFIXES
    | MARKDOWN_SUFFIXES
    | HTML_SUFFIXES
    | CSV_SUFFIXES
    | EXCEL_SUFFIXES
    | DOCX_SUFFIXES
    | PDF_SUFFIXES
    | JSON_SUFFIXES
    | CODE_SUFFIXES
)

HEADER_KEYS = {
    "SOURCE_URL",
    "SOURCE_URI",
    "SOURCE_FILE",
    "SOURCE_TYPE",
    "DOCUMENT_TITLE",
    "MANUAL_TYPE",
    "SECTION",
    "CLASS_CODE",
    "PAGE",
    "SHEET",
    "ROW_RANGE",
    "CHUNK_ID",
    "CHUNK_STRATEGY",
    "ALIASES",
    "SOURCE_HASH",
    "CONTENT_HASH",
    "LOCATOR",
}

EXPECTED_PROPERTY_SECTIONS = (
    "Appetite",
    "Applications",
    "Business Income",
    "CAT Exposed Property Wind and Hail Availability and Deductibles",
    "Causes of Loss",
    "Claims and Losses",
    "Class Occupancies and Minimum Premiums",
    "Class Specific Form Requirements",
    "Coastline Map",
    "Construction Type",
    "Deductibles",
    "GA Code Title 33, Chapter 32-11",
    "Inspections",
    "Limit Authority",
    "Optional Coverages",
    "Premium Modification",
    "Prohibited",
    "Solar Panels",
    "Triple Net Lease",
    "Vacant Buildings",
    "Valuation",
    "Wildfire Guide",
)

SECTION_ALIAS_HINTS = {
    "Appetite": (
        "property appetite",
        "market for package property",
        "monoline property",
        "building bpp business income tenant improvements",
    ),
    "Class Occupancies and Minimum Premiums": (
        "property class code",
        "minimum premium",
        "bpp only minimum premium",
    ),
    "Coastline Map": (
        "coastline map",
        "coastal map",
        "coast map",
        "gulf coast",
        "atlantic coast",
        "florida coastline",
    ),
    "Inspections": (
        "physical inspections",
        "inspection requirements",
        "tiv over 250000",
    ),
    "Optional Coverages": (
        "coverage options",
        "spoilage coverage",
        "ordinance or law",
        "property extension",
    ),
    "Solar Panels": (
        "photovoltaic panels",
        "pv panels",
        "building limit",
        "solar panels attached to a building",
    ),
    "Triple Net Lease": (
        "NNN lease",
        "triple net",
        "CP 12 19",
        "additional insured building owner",
    ),
    "Vacant Buildings": (
        "vacant building",
        "continuously vacant",
        "vacant more than 24 months",
        "structural renovations",
    ),
    "Wildfire Guide": (
        "risk meter",
        "hazardhub",
        "high wildfire",
        "very high wildfire",
        "active wildfire",
    ),
}

IMAGE_ONLY_SECTION_DESCRIPTIONS = {
    "Coastline Map": (
        "The Property Manual contains a Coastline Map image. The map highlights "
        "the U.S. coastline along the Gulf Coast, Florida, and the Atlantic Coast."
    ),
}

FORM_RE = re.compile(
    r"\b(?:CG|GL|BP|CP|IL|PR)\s*\d{2,4}(?:\s*\d{2,4}){0,3}\b",
    re.IGNORECASE,
)
CLASS_CODE_RE = re.compile(r"\b\d{4,}(?:-\d+)?\b")
URL_RE = re.compile(r"^https?://", re.IGNORECASE)


@dataclass(frozen=True)
class ContentBlock:
    heading: str
    text: str
    block_type: str = "section"
    page: int | None = None
    sheet: str | None = None
    row_start: int | None = None
    row_end: int | None = None
    block_index: int = 1
    char_start: int | None = None
    char_end: int | None = None

    @property
    def row_range(self) -> str | None:
        if self.row_start is None:
            return None
        if self.row_end is None or self.row_end == self.row_start:
            return str(self.row_start)
        return f"{self.row_start}-{self.row_end}"


@dataclass(frozen=True)
class SourceDocument:
    path: Path | None
    metadata: dict[str, str]
    body: str
    kind: str
    manual_type: str
    source_url: str
    source_uri: str
    class_code: str | None
    source_hash: str
    title: str
    parser: str = "deterministic"
    blocks: list[ContentBlock] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


@dataclass(frozen=True)
class Section:
    heading: str
    body: str
    level: int = 1


@dataclass
class ChunkRecord:
    source_path: str
    source_url: str
    manual_type: str
    section: str
    chunk_id: str
    source_hash: str
    content_hash: str
    content: str
    aliases: list[str] = field(default_factory=list)
    class_code: str | None = None
    source_uri: str = ""
    source_type: str = ""
    source_file: str = ""
    document_title: str = ""
    chunk_strategy: str = "hybrid"
    parser: str = "deterministic"
    page: int | None = None
    sheet: str | None = None
    row_range: str | None = None
    locator: str = ""
    warnings: list[str] = field(default_factory=list)

    def header_text(self) -> str:
        lines = [
            f"SOURCE_URL: {self.source_url}",
            f"SOURCE_URI: {self.source_uri or self.source_url}",
        ]
        if self.source_file:
            lines.append(f"SOURCE_FILE: {self.source_file}")
        lines.extend(
            [
                f"SOURCE_TYPE: {self.source_type}",
                f"DOCUMENT_TITLE: {self.document_title}",
                f"MANUAL_TYPE: {self.manual_type}",
                f"SECTION: {self.section}",
            ]
        )
        if self.class_code:
            lines.append(f"CLASS_CODE: {self.class_code}")
        if self.page is not None:
            lines.append(f"PAGE: {self.page}")
        if self.sheet:
            lines.append(f"SHEET: {self.sheet}")
        if self.row_range:
            lines.append(f"ROW_RANGE: {self.row_range}")
        lines.extend(
            [
                f"CHUNK_ID: {self.chunk_id}",
                f"CHUNK_STRATEGY: {self.chunk_strategy}",
                f"ALIASES: {'; '.join(self.aliases)}",
                f"SOURCE_HASH: {self.source_hash}",
                f"CONTENT_HASH: {self.content_hash}",
                f"LOCATOR: {self.locator}",
                "---",
                "",
            ]
        )
        return "\n".join(lines)

    def document_text(self) -> str:
        return f"{self.header_text()}{self.content.strip()}\n"


def _normalize_space(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def _slug(value: str, *, max_len: int = 90) -> str:
    value = re.sub(r"[_*`#]+", "", value).strip().lower()
    value = re.sub(r"[^a-z0-9]+", "-", value)
    value = value.strip("-")
    return value[:max_len].strip("-") or "chunk"


def _sha256(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8", errors="ignore")).hexdigest()


def _dedupe(values: Iterable[str]) -> list[str]:
    seen = set()
    result = []
    for value in values:
        clean = _normalize_space(str(value))
        if not clean:
            continue
        key = clean.lower()
        if key in seen:
            continue
        seen.add(key)
        result.append(clean)
    return result


def _clean_heading(value: str) -> str:
    value = re.sub(r"^#+\s*", "", value.strip())
    value = re.sub(r"^[-*o]\s+", "", value, flags=re.IGNORECASE)
    value = re.sub(r"[_*`]+", "", value)
    value = _normalize_space(value).strip(":-")
    if not value or value == "#" or len(value) > 140:
        return ""
    return value


def _clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    junk_patterns = (
        r"^\s*AI-generated content may be incorrect\.?\s*$",
        r"^\s*Was this helpful\??\s*$",
        r"^\s*Search\s+Previous\s+Next\s*$",
        r"^\s*Copy link\s*$",
    )
    lines = []
    for line in text.splitlines():
        if any(re.search(pattern, line, re.IGNORECASE) for pattern in junk_patterns):
            continue
        lines.append(line.rstrip())
    return "\n".join(lines).strip()


def _strip_existing_header(text: str) -> tuple[dict[str, str], str]:
    metadata: dict[str, str] = {}
    lines = text.splitlines()
    body_start = 0

    for idx, line in enumerate(lines[:60]):
        stripped = line.strip()
        if stripped.startswith("---"):
            body_start = idx + 1
            break
        match = re.match(r"^([A-Z][A-Z0-9_ ]{1,40}):\s*(.*)$", stripped)
        if not match:
            if idx == 0:
                break
            continue
        key = match.group(1).strip().replace(" ", "_")
        if key in HEADER_KEYS:
            metadata[key] = match.group(2).strip()
    else:
        body_start = 0 if not metadata else len(metadata)

    if metadata and body_start:
        return metadata, "\n".join(lines[body_start:]).strip()
    return metadata, text.strip()


def _is_url(value: str | Path) -> bool:
    return bool(URL_RE.match(str(value)))


def _source_uri_from_path(path: Path, metadata: dict[str, str]) -> str:
    return metadata.get("SOURCE_URI") or metadata.get("SOURCE_URL") or str(path)


def _source_url_from_path(path: Path, metadata: dict[str, str]) -> str:
    if metadata.get("SOURCE_URL"):
        return metadata["SOURCE_URL"]
    if path.stem.isdigit():
        return f"https://bindingauthority.coactionspecialty.com/manuals/{path.stem}.html"
    path_text = str(path).replace("\\", "/").lower()
    if "property" in path_text:
        return PROPERTY_MANUAL_URL
    if "guide" in path_text:
        return GL_GUIDE_MANUAL_URL
    return "N/A"


def _class_code_from_path_or_url(
    path: Path | None, source_url: str, metadata: dict[str, str]
) -> str | None:
    if metadata.get("CLASS_CODE"):
        return metadata["CLASS_CODE"]
    if path and path.stem.isdigit():
        return path.stem
    match = re.search(r"/(\d{4,})\.(?:html|md)(?:$|[?#])", source_url)
    return match.group(1) if match else None


def detect_document_kind(
    path: Path | None, metadata: dict[str, str], body: str
) -> tuple[str, str]:
    """Return (source kind, manual/content type)."""
    manual_type = metadata.get("MANUAL_TYPE", "")
    source_url = metadata.get("SOURCE_URL", "")
    path_text = str(path or "")
    haystack = f"{manual_type} {source_url} {path_text} {body[:800]}".lower().replace("\\", "/")

    if "internal" in haystack or "binding authority and brokerage light internal" in haystack:
        return "internal_guidelines", "Internal Guidelines"
    if metadata.get("CLASS_CODE") or re.search(r"/manuals/\d{4,}\.html", haystack):
        return "gl_class_code", "General Liability"
    if "property" in manual_type.lower() or "property.html" in haystack or "/property/" in haystack:
        return "property_manual", "Property"
    if (
        "guide.html" in haystack
        or "/guide_" in haystack
        or "additional insured" in body[:1000].lower()
    ):
        return "gl_guide_manual", "General Liability Guide"

    suffix = (path.suffix.lower() if path else "").strip()
    if suffix in PDF_SUFFIXES:
        return "pdf", manual_type or "Document"
    if suffix in DOCX_SUFFIXES:
        return "word", manual_type or "Document"
    if suffix in EXCEL_SUFFIXES or suffix in CSV_SUFFIXES:
        return "spreadsheet", manual_type or "Spreadsheet"
    if suffix in HTML_SUFFIXES or _is_url(source_url):
        return "website", manual_type or "Website"
    if suffix in CODE_SUFFIXES:
        return "code", manual_type or "Code"
    return "unknown", manual_type or "Document"


def _title_from_body(body: str, fallback: str) -> str:
    for line in body.splitlines():
        match = re.match(r"^#\s+(.+)$", line.strip())
        if match:
            return _clean_heading(match.group(1)) or fallback
    return fallback


def _file_title(path: Path | None, source_uri: str) -> str:
    if path:
        return path.stem.replace("_", " ").replace("-", " ").strip() or path.name
    parsed = urlparse(source_uri)
    name = Path(parsed.path).stem if parsed.path else parsed.netloc
    return name.replace("_", " ").replace("-", " ").strip() or source_uri


def _text_blocks_from_markdown(body: str, default_heading: str) -> list[ContentBlock]:
    sections = split_markdown_sections_from_text(body, default_heading)
    blocks = []
    running = 0
    for idx, section in enumerate(sections, 1):
        text = section.body.strip()
        char_start = body.find(text, running) if text else running
        char_end = char_start + len(text) if char_start >= 0 else None
        running = char_end or running
        blocks.append(
            ContentBlock(
                heading=section.heading,
                text=text,
                block_index=idx,
                char_start=char_start if char_start >= 0 else None,
                char_end=char_end,
            )
        )
    return blocks


def split_markdown_sections_from_text(body: str, default_heading: str = "Manual Section") -> list[Section]:
    sections: list[Section] = []
    current_heading = _clean_heading(default_heading) or "Manual Section"
    current_level = 1
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_lines
        section_body = "\n".join(current_lines).strip()
        heading = _clean_heading(current_heading)
        if heading and not section_body:
            section_body = IMAGE_ONLY_SECTION_DESCRIPTIONS.get(heading, "")
        if heading and section_body:
            sections.append(Section(heading=heading, body=section_body, level=current_level))
        current_lines = []

    for line in body.splitlines():
        match = re.match(r"^(#{1,6})\s*(.+?)\s*$", line)
        if match:
            heading = _clean_heading(match.group(2))
            if heading:
                flush()
                current_heading = heading
                current_level = len(match.group(1))
                current_lines = []
                continue
        current_lines.append(line)

    flush()
    if not sections and body.strip():
        sections.append(Section(heading=current_heading, body=body.strip()))
    return [section for section in sections if not _is_toc_only(section)]


def split_markdown_sections(doc: SourceDocument) -> list[Section]:
    """Compatibility helper used by tests and older scripts."""
    default_heading = doc.metadata.get("SECTION") or (
        f"Class Code {doc.class_code}" if doc.class_code else "Manual Section"
    )
    return split_markdown_sections_from_text(doc.body, default_heading)


def _is_toc_only(section: Section) -> bool:
    lines = [line.strip() for line in section.body.splitlines() if line.strip()]
    if not lines:
        return True
    bullet_like = sum(1 for line in lines if re.match(r"^[-*o]\s+\w", line, re.IGNORECASE))
    has_sentence = any(
        re.search(r"\b(is|are|must|required|include|available|prohibited)\b", line, re.I)
        for line in lines
    )
    return len(lines) >= 5 and bullet_like / max(len(lines), 1) > 0.8 and not has_sentence


class _SimpleHTMLToMarkdown(HTMLParser):
    """Small stdlib fallback when BeautifulSoup is not installed."""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._parts: list[str] = []
        self._tag_stack: list[str] = []
        self._skip_depth = 0
        self._cell: list[str] | None = None
        self._row: list[str] = []
        self._rows: list[list[str]] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag in {"script", "style", "nav", "footer", "button", "form"}:
            self._skip_depth += 1
            return
        self._tag_stack.append(tag)
        if tag in {"h1", "h2", "h3", "h4", "h5", "h6", "p", "div", "section", "article"}:
            self._parts.append("\n")
        elif tag == "li":
            self._parts.append("\n- ")
        elif tag == "tr":
            self._row = []
        elif tag in {"td", "th"}:
            self._cell = []

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if self._skip_depth:
            self._skip_depth -= 1
            return
        if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = int(tag[1])
            line = self._take_current_line()
            if line:
                self._parts.append(f"\n{'#' * level} {line}\n")
        elif tag in {"td", "th"} and self._cell is not None:
            self._row.append(_normalize_space(" ".join(self._cell)))
            self._cell = None
        elif tag == "tr" and self._row:
            self._rows.append(self._row)
            self._row = []
        elif tag == "table" and self._rows:
            self._parts.append("\n" + _markdown_table(self._rows) + "\n")
            self._rows = []
        elif tag in {"p", "div", "section", "article", "li"}:
            self._parts.append("\n")
        if self._tag_stack:
            self._tag_stack.pop()

    def handle_data(self, data: str) -> None:
        if self._skip_depth:
            return
        value = html.unescape(data)
        if self._cell is not None:
            self._cell.append(value)
        else:
            self._parts.append(value)

    def _take_current_line(self) -> str:
        text = "".join(self._parts)
        lines = text.splitlines()
        if not lines:
            return ""
        current = _normalize_space(lines[-1])
        self._parts = ["\n".join(lines[:-1])]
        return current

    def markdown(self) -> str:
        return _clean_text("\n".join(self._parts))


def _markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in normalized[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def html_to_clean_markdown(raw_html: str, *, source_url: str = "") -> str:
    """Clean website/local HTML into Markdown-like text with tables preserved."""
    try:
        from bs4 import BeautifulSoup, Comment

        soup = BeautifulSoup(raw_html, "html.parser")
        for node in soup(["script", "style", "noscript", "svg", "form", "button"]):
            node.decompose()
        for comment in soup.find_all(string=lambda value: isinstance(value, Comment)):
            comment.extract()

        root = (
            soup.find("main")
            or soup.find("article")
            or soup.find(attrs={"role": "main"})
            or soup.body
            or soup
        )
        parts: list[str] = []
        for element in root.descendants:
            name = getattr(element, "name", None)
            if name in {"nav", "footer", "aside"}:
                continue
            if name and re.fullmatch(r"h[1-6]", name):
                level = int(name[1])
                heading = _normalize_space(element.get_text(" ", strip=True))
                if heading:
                    parts.append(f"\n{'#' * level} {heading}\n")
            elif name == "p":
                text = _normalize_space(element.get_text(" ", strip=True))
                if text:
                    parts.append(f"{text}\n")
            elif name == "li":
                text = _normalize_space(element.get_text(" ", strip=True))
                if text:
                    parts.append(f"- {text}\n")
            elif name == "table":
                rows: list[list[str]] = []
                for tr in element.find_all("tr"):
                    cells = [
                        _normalize_space(cell.get_text(" ", strip=True))
                        for cell in tr.find_all(["th", "td"])
                    ]
                    if any(cells):
                        rows.append(cells)
                table = _markdown_table(rows)
                if table:
                    parts.append(table + "\n")
        markdown = "\n".join(parts)
    except Exception:
        parser = _SimpleHTMLToMarkdown()
        parser.feed(raw_html)
        markdown = parser.markdown()

    if source_url and not re.search(r"^SOURCE_URL:", markdown, re.MULTILINE):
        markdown = f"SOURCE_URL: {source_url}\n---\n\n{markdown}"
    return _clean_text(markdown)


def _fetch_url(url: str, *, timeout: int = 30) -> str:
    request = Request(url, headers={"User-Agent": "ProjectVegaChunker/1.0"})
    with urlopen(request, timeout=timeout) as response:  # noqa: S310 - operator-provided URL
        encoding = response.headers.get_content_charset() or "utf-8"
        return response.read().decode(encoding, errors="replace")


def _parse_markdown_file(path: Path) -> SourceDocument:
    text = _clean_text(path.read_text(encoding="utf-8", errors="replace"))
    metadata, body = _strip_existing_header(text)
    source_url = _source_url_from_path(path, metadata)
    source_uri = _source_uri_from_path(path, metadata)
    class_code = _class_code_from_path_or_url(path, source_url, metadata)
    metadata.setdefault("SOURCE_URL", source_url)
    metadata.setdefault("SOURCE_URI", source_uri)
    if class_code:
        metadata.setdefault("CLASS_CODE", class_code)
    kind, manual_type = detect_document_kind(path, metadata, body)
    metadata.setdefault("MANUAL_TYPE", manual_type)
    default_heading = metadata.get("SECTION") or (
        f"Class Code {class_code}" if class_code else "Manual Section"
    )
    title = metadata.get("DOCUMENT_TITLE") or _title_from_body(body, _file_title(path, source_uri))
    blocks = _text_blocks_from_markdown(body, default_heading)
    return SourceDocument(
        path=path,
        metadata=metadata,
        body=body,
        kind=kind,
        manual_type=manual_type,
        source_url=source_url,
        source_uri=source_uri,
        class_code=class_code,
        source_hash=_sha256(body),
        title=title,
        blocks=blocks,
    )


def _parse_text_file(path: Path) -> SourceDocument:
    body = _clean_text(path.read_text(encoding="utf-8", errors="replace"))
    metadata: dict[str, str] = {"SOURCE_URI": str(path)}
    source_url = _source_url_from_path(path, metadata)
    kind, manual_type = detect_document_kind(path, metadata, body)
    title = _file_title(path, str(path))
    block = ContentBlock(heading=title, text=body, block_index=1, char_start=0, char_end=len(body))
    return SourceDocument(
        path=path,
        metadata=metadata,
        body=body,
        kind=kind,
        manual_type=manual_type,
        source_url=source_url,
        source_uri=str(path),
        class_code=_class_code_from_path_or_url(path, source_url, metadata),
        source_hash=_sha256(body),
        title=title,
        blocks=[block],
    )


def _parse_html_file(path: Path) -> SourceDocument:
    raw = path.read_text(encoding="utf-8", errors="replace")
    source_url = _source_url_from_path(path, {})
    markdown = html_to_clean_markdown(raw, source_url=source_url)
    tmp = path.with_suffix(".md")
    metadata, body = _strip_existing_header(markdown)
    metadata.setdefault("SOURCE_URL", source_url)
    metadata.setdefault("SOURCE_URI", str(path))
    kind, manual_type = detect_document_kind(tmp, metadata, body)
    title = _title_from_body(body, _file_title(path, source_url))
    blocks = _text_blocks_from_markdown(body, title)
    return SourceDocument(
        path=path,
        metadata=metadata,
        body=body,
        kind=kind,
        manual_type=manual_type,
        source_url=source_url,
        source_uri=str(path),
        class_code=_class_code_from_path_or_url(path, source_url, metadata),
        source_hash=_sha256(body),
        title=title,
        blocks=blocks,
    )


def _parse_website(url: str) -> SourceDocument:
    raw = _fetch_url(url)
    markdown = html_to_clean_markdown(raw, source_url=url)
    metadata, body = _strip_existing_header(markdown)
    metadata.setdefault("SOURCE_URL", url)
    metadata.setdefault("SOURCE_URI", url)
    kind, manual_type = detect_document_kind(None, metadata, body)
    title = _title_from_body(body, _file_title(None, url))
    blocks = _text_blocks_from_markdown(body, title)
    return SourceDocument(
        path=None,
        metadata=metadata,
        body=body,
        kind=kind,
        manual_type=manual_type,
        source_url=url,
        source_uri=url,
        class_code=_class_code_from_path_or_url(None, url, metadata),
        source_hash=_sha256(body),
        title=title,
        parser="website",
        blocks=blocks,
    )


def _parse_csv_file(path: Path) -> SourceDocument:
    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as handle:
        rows = list(csv.reader(handle, delimiter=delimiter))
    title = _file_title(path, str(path))
    if not rows:
        body = ""
        blocks: list[ContentBlock] = []
    else:
        header = rows[0]
        body_parts = [_markdown_table(rows)]
        blocks = []
        group_size = 25
        for start in range(1, len(rows), group_size):
            group = rows[start : start + group_size]
            table = _markdown_table([header, *group])
            blocks.append(
                ContentBlock(
                    heading=title,
                    text=table,
                    block_type="table",
                    sheet=path.stem,
                    row_start=start + 1,
                    row_end=start + len(group),
                    block_index=len(blocks) + 1,
                )
            )
        body = "\n\n".join(body_parts)
    metadata = {"SOURCE_URI": str(path)}
    kind, manual_type = detect_document_kind(path, metadata, body)
    return SourceDocument(
        path=path,
        metadata=metadata,
        body=body,
        kind=kind,
        manual_type=manual_type,
        source_url=_source_url_from_path(path, metadata),
        source_uri=str(path),
        class_code=None,
        source_hash=_sha256(body),
        title=title,
        blocks=blocks,
    )


def _parse_excel_file(path: Path) -> SourceDocument:
    try:
        import openpyxl
    except Exception as exc:
        return _parse_with_agno_or_error(path, f"openpyxl_not_installed: {exc}")

    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    blocks: list[ContentBlock] = []
    body_parts: list[str] = []
    for sheet in workbook.worksheets:
        rows = [
            ["" if value is None else str(value) for value in row]
            for row in sheet.iter_rows(values_only=True)
            if any(value is not None and str(value).strip() for value in row)
        ]
        if not rows:
            continue
        sheet_heading = f"{path.stem} - {sheet.title}"
        body_parts.append(f"# {sheet_heading}\n\n{_markdown_table(rows)}")
        header = rows[0]
        for start in range(1, len(rows), 25):
            group = rows[start : start + 25]
            blocks.append(
                ContentBlock(
                    heading=sheet_heading,
                    text=_markdown_table([header, *group]),
                    block_type="table",
                    sheet=sheet.title,
                    row_start=start + 1,
                    row_end=start + len(group),
                    block_index=len(blocks) + 1,
                )
            )
    body = "\n\n".join(body_parts)
    metadata = {"SOURCE_URI": str(path)}
    kind, manual_type = detect_document_kind(path, metadata, body)
    return SourceDocument(
        path=path,
        metadata=metadata,
        body=body,
        kind=kind,
        manual_type=manual_type,
        source_url=_source_url_from_path(path, metadata),
        source_uri=str(path),
        class_code=None,
        source_hash=_sha256(body),
        title=_file_title(path, str(path)),
        blocks=blocks,
        parser="openpyxl",
    )


def _docx_paragraphs_with_stdlib(path: Path) -> list[tuple[str, str]]:
    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    paragraphs: list[tuple[str, str]] = []
    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml")
    root = ET.fromstring(document_xml)
    for paragraph in root.findall(".//w:p", namespaces):
        style = ""
        style_el = paragraph.find(".//w:pStyle", namespaces)
        if style_el is not None:
            style = style_el.attrib.get(f"{{{namespaces['w']}}}val", "")
        text = "".join(node.text or "" for node in paragraph.findall(".//w:t", namespaces))
        text = _normalize_space(text)
        if text:
            paragraphs.append((style, text))
    return paragraphs


def _parse_docx_file(path: Path) -> SourceDocument:
    body_lines: list[str] = []
    parser = "docx-stdlib"
    try:
        from docx import Document as DocxDocument

        parser = "python-docx"
        docx = DocxDocument(path)
        for paragraph in docx.paragraphs:
            text = _normalize_space(paragraph.text)
            if not text:
                continue
            style = (paragraph.style.name if paragraph.style else "").lower()
            if style.startswith("heading"):
                level_match = re.search(r"(\d+)", style)
                level = int(level_match.group(1)) if level_match else 2
                body_lines.append(f"{'#' * max(1, min(level, 6))} {text}")
            else:
                body_lines.append(text)
        for table in docx.tables:
            rows = [[_normalize_space(cell.text) for cell in row.cells] for row in table.rows]
            if rows:
                body_lines.append(_markdown_table(rows))
    except Exception:
        paragraphs = _docx_paragraphs_with_stdlib(path)
        for style, text in paragraphs:
            if style.lower().startswith("heading"):
                level_match = re.search(r"(\d+)", style)
                level = int(level_match.group(1)) if level_match else 2
                body_lines.append(f"{'#' * max(1, min(level, 6))} {text}")
            else:
                body_lines.append(text)

    body = _clean_text("\n\n".join(body_lines))
    metadata = {"SOURCE_URI": str(path)}
    kind, manual_type = detect_document_kind(path, metadata, body)
    title = _title_from_body(body, _file_title(path, str(path)))
    blocks = _text_blocks_from_markdown(body, title)
    return SourceDocument(
        path=path,
        metadata=metadata,
        body=body,
        kind=kind,
        manual_type=manual_type,
        source_url=_source_url_from_path(path, metadata),
        source_uri=str(path),
        class_code=_class_code_from_path_or_url(path, "N/A", metadata),
        source_hash=_sha256(body),
        title=title,
        parser=parser,
        blocks=blocks,
    )


def _parse_pdf_file(path: Path) -> SourceDocument:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        return _parse_with_agno_or_error(path, f"pypdf_not_installed: {exc}")

    reader = PdfReader(str(path))
    blocks: list[ContentBlock] = []
    body_parts: list[str] = []
    for idx, page in enumerate(reader.pages, 1):
        text = _clean_text(page.extract_text() or "")
        if not text:
            continue
        heading = f"{path.stem} - Page {idx}"
        body_parts.append(f"# {heading}\n\n{text}")
        blocks.append(
            ContentBlock(
                heading=heading,
                text=text,
                block_type="page",
                page=idx,
                block_index=len(blocks) + 1,
            )
        )
    body = "\n\n".join(body_parts)
    metadata = {"SOURCE_URI": str(path)}
    kind, manual_type = detect_document_kind(path, metadata, body)
    return SourceDocument(
        path=path,
        metadata=metadata,
        body=body,
        kind=kind,
        manual_type=manual_type,
        source_url=_source_url_from_path(path, metadata),
        source_uri=str(path),
        class_code=None,
        source_hash=_sha256(body),
        title=_file_title(path, str(path)),
        parser="pypdf",
        blocks=blocks,
    )


def _parse_json_file(path: Path) -> SourceDocument:
    data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    title = _file_title(path, str(path))
    blocks: list[ContentBlock] = []
    if isinstance(data, list):
        for idx, item in enumerate(data, 1):
            blocks.append(
                ContentBlock(
                    heading=f"{title} item {idx}",
                    text=json.dumps(item, indent=2, ensure_ascii=False),
                    block_type="json",
                    row_start=idx,
                    row_end=idx,
                    block_index=idx,
                )
            )
    elif isinstance(data, dict):
        for idx, (key, value) in enumerate(data.items(), 1):
            blocks.append(
                ContentBlock(
                    heading=f"{title} - {key}",
                    text=json.dumps(value, indent=2, ensure_ascii=False),
                    block_type="json",
                    block_index=idx,
                )
            )
    else:
        blocks.append(ContentBlock(heading=title, text=json.dumps(data), block_type="json"))
    body = "\n\n".join(f"# {block.heading}\n\n{block.text}" for block in blocks)
    metadata = {"SOURCE_URI": str(path)}
    kind, manual_type = detect_document_kind(path, metadata, body)
    return SourceDocument(
        path=path,
        metadata=metadata,
        body=body,
        kind=kind,
        manual_type=manual_type,
        source_url=_source_url_from_path(path, metadata),
        source_uri=str(path),
        class_code=None,
        source_hash=_sha256(body),
        title=title,
        blocks=blocks,
    )


def _parse_with_agno_or_error(path: Path, warning: str) -> SourceDocument:
    agno_doc = _parse_with_agno_reader(str(path), path=path)
    if agno_doc:
        agno_doc.warnings.append(warning)
        return agno_doc
    title = _file_title(path, str(path))
    body = f"Unable to parse {path.name}. {warning}"
    metadata = {"SOURCE_URI": str(path)}
    return SourceDocument(
        path=path,
        metadata=metadata,
        body=body,
        kind="unsupported",
        manual_type="Document",
        source_url=_source_url_from_path(path, metadata),
        source_uri=str(path),
        class_code=None,
        source_hash=_sha256(body),
        title=title,
        blocks=[ContentBlock(heading=title, text=body, block_index=1)],
        warnings=[warning],
    )


def _parse_with_agno_reader(source: str, *, path: Path | None = None) -> SourceDocument | None:
    """Best-effort Agno ReaderFactory parser. Returns None if Agno is unavailable."""
    try:
        from agno.knowledge.reader.reader_factory import ReaderFactory
    except Exception:
        return None

    try:
        if _is_url(source):
            reader = ReaderFactory.get_reader_for_url(source)
            documents = reader.read(source)
        else:
            reader = ReaderFactory.get_reader_for_extension(Path(source).suffix)
            documents = reader.read(source)
    except Exception:
        return None

    if not documents:
        return None

    blocks: list[ContentBlock] = []
    body_parts: list[str] = []
    source_url = source if _is_url(source) else _source_url_from_path(path or Path(source), {})
    source_uri = source
    title = _file_title(path, source)
    for idx, doc in enumerate(documents, 1):
        content = _clean_text(str(getattr(doc, "content", "") or ""))
        if not content:
            continue
        meta = getattr(doc, "meta_data", {}) or {}
        heading = _clean_heading(meta.get("section") or meta.get("name") or title) or title
        page = _safe_int(meta.get("page") or meta.get("page_number"))
        sheet = str(meta.get("sheet") or "") or None
        blocks.append(
            ContentBlock(
                heading=heading,
                text=content,
                block_type="agno",
                page=page,
                sheet=sheet,
                block_index=idx,
            )
        )
        body_parts.append(f"# {heading}\n\n{content}")
    body = "\n\n".join(body_parts)
    metadata = {"SOURCE_URL": source_url, "SOURCE_URI": source_uri}
    kind, manual_type = detect_document_kind(path, metadata, body)
    return SourceDocument(
        path=path,
        metadata=metadata,
        body=body,
        kind=kind,
        manual_type=manual_type,
        source_url=source_url,
        source_uri=source_uri,
        class_code=_class_code_from_path_or_url(path, source_url, metadata),
        source_hash=_sha256(body),
        title=title,
        parser="agno-reader",
        blocks=blocks,
    )


def _safe_int(value: Any) -> int | None:
    try:
        return int(value)
    except Exception:
        return None


def load_source_document(path: Path) -> SourceDocument:
    suffix = path.suffix.lower()
    if suffix in MARKDOWN_SUFFIXES:
        return _parse_markdown_file(path)
    if suffix in HTML_SUFFIXES:
        return _parse_html_file(path)
    if suffix in TEXT_SUFFIXES or suffix in CODE_SUFFIXES:
        return _parse_text_file(path)
    if suffix in CSV_SUFFIXES:
        return _parse_csv_file(path)
    if suffix in EXCEL_SUFFIXES:
        return _parse_excel_file(path)
    if suffix in DOCX_SUFFIXES:
        return _parse_docx_file(path)
    if suffix in PDF_SUFFIXES:
        return _parse_pdf_file(path)
    if suffix in JSON_SUFFIXES:
        return _parse_json_file(path)

    agno_doc = _parse_with_agno_reader(str(path), path=path)
    if agno_doc:
        return agno_doc

    return _parse_text_file(path)


def load_source_reference(source: str | Path, *, prefer_agno_reader: bool = False) -> list[SourceDocument]:
    if _is_url(source):
        source_text = str(source)
        if prefer_agno_reader:
            agno_doc = _parse_with_agno_reader(source_text)
            if agno_doc:
                return [agno_doc]
        return [_parse_website(source_text)]

    path = Path(source)
    if path.is_dir():
        return [load_source_document(candidate) for candidate in collect_source_files([path])]
    if prefer_agno_reader:
        agno_doc = _parse_with_agno_reader(str(path), path=path)
        if agno_doc:
            return [agno_doc]
    return [load_source_document(path)]


def _extract_forms(text: str) -> list[str]:
    return _dedupe(_normalize_space(match.group(0).upper()) for match in FORM_RE.finditer(text))


def _form_aliases(form: str) -> list[str]:
    cleaned = _normalize_space(form.upper())
    compact = re.sub(r"\s+", "", cleaned)
    aliases = [cleaned, compact]
    match = re.match(r"^([A-Z]{2})\s*(\d{4})(?:\s+(\d{2,4}))?$", cleaned)
    if match:
        prefix, number, suffix = match.groups()
        aliases.append(f"{prefix} {number[:2]} {number[2:]}")
        if suffix:
            aliases.append(f"{prefix}{number}{suffix}")
    return _dedupe(aliases)


def _class_title_aliases(doc: SourceDocument) -> list[str]:
    if not doc.class_code:
        return []
    aliases = [doc.class_code, f"Class Code {doc.class_code}"]
    pattern = rf"\b{re.escape(doc.class_code)}\s*[-:]\s*([^\n#]+)"
    match = re.search(pattern, doc.body)
    if not match:
        return aliases
    title = _normalize_space(re.sub(r"[_*`]+", "", match.group(1))).strip(" -")
    if title:
        aliases.append(title)
        aliases.append(re.sub(r"\bMFG\b", "Manufacturing", title, flags=re.IGNORECASE))
    return aliases


def _section_keyword_aliases(section: str, text: str) -> list[str]:
    if section not in {"Prohibited", "Submit", "Coverage Options", "Class-Specific Forms"}:
        return []
    aliases: list[str] = []
    for line in text.splitlines():
        clean = re.sub(r"^#+\s*", "", line.strip())
        clean = re.sub(r"^[-*o]\s+", "", clean, flags=re.IGNORECASE)
        clean = re.sub(r"^\|\s*", "", clean)
        clean = re.sub(r"\s*\|.*$", "", clean)
        clean = re.sub(r"[_*`]+", "", clean)
        clean = _normalize_space(clean).strip(" :-")
        if 3 <= len(clean) <= 90 and re.search(r"[A-Za-z]", clean):
            aliases.append(clean)
        if len(aliases) >= 20:
            break
    return aliases


def _section_aliases(section: str, text: str, doc: SourceDocument | None = None) -> list[str]:
    aliases = [section]
    if doc:
        aliases.extend(_class_title_aliases(doc))
    aliases.extend(SECTION_ALIAS_HINTS.get(section, ()))
    aliases.extend(_section_keyword_aliases(section, text))
    for form in _extract_forms(text):
        aliases.extend(_form_aliases(form))
    return _dedupe(aliases)


def _section_text(block: ContentBlock) -> str:
    return f"# {block.heading}\n\n{block.text.strip()}".strip()


def choose_chunking_strategy(doc: SourceDocument, requested: str = "auto") -> str:
    requested = requested.lower()
    if requested != "auto":
        return requested
    if doc.kind in {"spreadsheet"}:
        return "table"
    if doc.kind in {"property_manual", "gl_guide_manual", "gl_class_code", "website", "word"}:
        return "hybrid"
    if doc.kind == "pdf":
        return "page_hybrid"
    if doc.kind == "code":
        return "recursive"
    return "recursive"


def _paragraph_chunks(text: str, max_chars: int) -> list[str]:
    if len(text) <= max_chars:
        return [text.strip()]

    paragraphs = re.split(r"\n\s*\n", text)
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        if current and len(current) + len(paragraph) + 2 > max_chars:
            chunks.append(current.strip())
            current = paragraph
        elif not current:
            current = paragraph
        else:
            current = f"{current}\n\n{paragraph}"
    if current.strip():
        chunks.append(current.strip())

    split_chunks: list[str] = []
    for chunk in chunks:
        if len(chunk) <= max_chars:
            split_chunks.append(chunk)
            continue
        split_chunks.extend(_sentence_chunks(chunk, max_chars))
    return [chunk for chunk in split_chunks if chunk]


def _sentence_chunks(text: str, max_chars: int) -> list[str]:
    sentences = re.split(r"(?<=[.!?])\s+", text)
    chunks: list[str] = []
    current = ""
    for sentence in sentences:
        if current and len(current) + len(sentence) + 1 > max_chars:
            chunks.append(current.strip())
            current = sentence
        elif not current:
            current = sentence
        else:
            current = f"{current} {sentence}"
    if current.strip():
        chunks.append(current.strip())
    final: list[str] = []
    for chunk in chunks:
        if len(chunk) <= max_chars:
            final.append(chunk)
        else:
            for start in range(0, len(chunk), max_chars):
                final.append(chunk[start : start + max_chars].strip())
    return final


def _agno_split_text(text: str, strategy: str, max_chars: int) -> list[str] | None:
    """Use Agno as a best-effort boundary selector. Source text is not rewritten."""
    if strategy not in {"agentic", "semantic"}:
        return None
    try:
        if strategy == "agentic":
            from agno.knowledge.chunking.agentic import AgenticChunking as Chunker

            chunker = Chunker(max_chunk_size=max_chars)
        else:
            try:
                from agno.knowledge.chunking.semantic import SemanticChunking as Chunker
            except Exception:
                from agno.knowledge.chunking.semantic_chunking import SemanticChunking as Chunker

            chunker = Chunker(chunk_size=max_chars)
        try:
            from agno.knowledge.document import Document
        except Exception:
            from agno.knowledge.document.base import Document
    except Exception:
        return None

    try:
        docs = chunker.chunk(Document(content=text, name="project-vega-section"))
    except Exception:
        return None

    chunks = [str(getattr(doc, "content", "")).strip() for doc in docs if getattr(doc, "content", "")]
    if not chunks:
        return None
    original_compact = re.sub(r"\s+", "", text)
    combined_compact = re.sub(r"\s+", "", "".join(chunks))
    # Guard against LLM/chunker rewrites. If content no longer resembles the
    # source, fall back to deterministic splitting.
    if combined_compact[:80] not in original_compact:
        return None
    return chunks


def split_long_text(
    text: str,
    max_chars: int,
    *,
    strategy: str,
    use_agno: bool,
) -> list[str]:
    if len(text) <= max_chars:
        return [text.strip()]
    if use_agno and strategy in {"agentic", "semantic"}:
        chunks = _agno_split_text(text, strategy, max_chars)
        if chunks:
            return chunks
    if use_agno and strategy in {"hybrid", "page_hybrid"}:
        chunks = _agno_split_text(text, "agentic", max_chars)
        if chunks:
            return chunks
    return _paragraph_chunks(text, max_chars)


def _first_table_header(lines: list[str], row_index: int) -> list[str]:
    start = row_index
    while start > 0 and lines[start - 1].lstrip().startswith("|"):
        start -= 1
    table_lines = [line for line in lines[start : row_index + 1] if line.lstrip().startswith("|")]
    header = table_lines[:2]
    if len(header) == 2 and re.search(r"\|\s*-{2,}", header[1]):
        return header
    return []


def _table_row_text(lines: list[str], row_index: int) -> str:
    row_lines = [lines[row_index]]
    next_index = row_index + 1
    while next_index < len(lines):
        line = lines[next_index]
        if line.lstrip().startswith("|"):
            break
        if not line.strip():
            break
        row_lines.append(line)
        if len(row_lines) >= 8:
            break
        next_index += 1
    return "\n".join(row_lines).strip()


def _looks_like_entity_row(row_text: str) -> bool:
    if re.search(r"\|\s*-{2,}", row_text):
        return False
    if FORM_RE.search(row_text):
        return True
    if re.search(r"^\|\s*\d{3,5}(?:-\d+)?\s*\|", row_text):
        return True
    if re.search(r"\b(Spoilage Coverage|Solar Panels|Triple Net|Vacant Building)\b", row_text, re.I):
        return True
    return False


def extract_table_entity_chunks(section: Section | ContentBlock) -> list[tuple[str, list[str]]]:
    heading = section.heading
    body = section.body if isinstance(section, Section) else section.text
    lines = body.splitlines()
    entities: list[tuple[str, list[str]]] = []
    for idx, line in enumerate(lines):
        if not line.lstrip().startswith("|"):
            continue
        row_text = _table_row_text(lines, idx)
        if not _looks_like_entity_row(row_text):
            continue
        header = _first_table_header(lines, idx)
        chunk_lines = [f"# {heading}", "", "Table excerpt:"]
        if header:
            chunk_lines.extend(header)
        chunk_lines.append(row_text)
        aliases = []
        aliases.extend(_extract_forms(row_text))
        if code_match := re.match(r"^\|\s*(\d{3,5}(?:-\d+)?)\s*\|", row_text):
            aliases.append(code_match.group(1))
        first_cell = row_text.split("|")[1].strip() if "|" in row_text else ""
        aliases.append(first_cell)
        entities.append(("\n".join(chunk_lines).strip(), aliases))
    return entities


def _chunk_id(doc: SourceDocument, section: str, ordinal: int, content: str) -> str:
    base = _slug(f"{doc.manual_type}-{doc.class_code or ''}-{section}")
    digest = _sha256(f"{doc.source_hash}:{ordinal}:{content}")[:10]
    return f"{base}-{ordinal:03d}-{digest}"


def _locator(doc: SourceDocument, block: ContentBlock, piece_index: int = 1) -> str:
    parts = [f"block={block.block_index}", f"piece={piece_index}", f"parser={doc.parser}"]
    if block.page is not None:
        parts.append(f"page={block.page}")
    if block.sheet:
        parts.append(f"sheet={block.sheet}")
    if block.row_range:
        parts.append(f"rows={block.row_range}")
    if block.char_start is not None and block.char_end is not None:
        parts.append(f"chars={block.char_start}-{block.char_end}")
    return "; ".join(parts)


def build_chunks_for_document(
    doc: SourceDocument,
    *,
    max_chars: int = DEFAULT_MAX_CHARS,
    strategy: str = "auto",
    use_agno: bool = False,
) -> list[ChunkRecord]:
    selected_strategy = choose_chunking_strategy(doc, strategy)
    chunks: list[ChunkRecord] = []
    ordinal = 1
    blocks = doc.blocks or _text_blocks_from_markdown(doc.body, doc.title)

    for block in blocks:
        block_content = _section_text(block)
        split_strategy = selected_strategy
        if selected_strategy == "table":
            split_strategy = "recursive"
        pieces = split_long_text(
            block_content,
            max_chars,
            strategy=split_strategy,
            use_agno=use_agno,
        )
        for piece_index, piece in enumerate(pieces, 1):
            if len(piece.strip()) < MIN_BODY_CHARS:
                continue
            chunks.append(
                _make_chunk(
                    doc=doc,
                    block=block,
                    section=block.heading,
                    content=piece,
                    aliases=_section_aliases(block.heading, piece, doc),
                    ordinal=ordinal,
                    strategy=selected_strategy,
                    locator=_locator(doc, block, piece_index),
                )
            )
            ordinal += 1

        if selected_strategy in {"hybrid", "page_hybrid", "table"}:
            for entity_content, entity_aliases in extract_table_entity_chunks(block):
                entity_aliases = _dedupe(
                    [*_section_aliases(block.heading, entity_content, doc), *entity_aliases]
                )
                chunks.append(
                    _make_chunk(
                        doc=doc,
                        block=block,
                        section=block.heading,
                        content=entity_content,
                        aliases=entity_aliases,
                        ordinal=ordinal,
                        strategy=f"{selected_strategy}+table_row",
                        locator=_locator(doc, block, 1),
                    )
                )
                ordinal += 1
    return chunks


def build_chunks_for_source(
    path: Path,
    *,
    max_chars: int = DEFAULT_MAX_CHARS,
    use_agno: bool = False,
    strategy: str = "auto",
) -> list[ChunkRecord]:
    return build_chunks_for_document(
        load_source_document(path),
        max_chars=max_chars,
        strategy=strategy,
        use_agno=use_agno,
    )


def _make_chunk(
    *,
    doc: SourceDocument,
    block: ContentBlock,
    section: str,
    content: str,
    aliases: list[str],
    ordinal: int,
    strategy: str,
    locator: str,
) -> ChunkRecord:
    content = content.strip()
    content_hash = _sha256(f"{doc.source_url}:{doc.class_code or ''}:{section}:{content}")
    return ChunkRecord(
        source_path=str(doc.path or doc.source_uri),
        source_url=doc.source_url,
        source_uri=doc.source_uri,
        source_type=doc.kind,
        source_file=str(doc.path) if doc.path else "",
        document_title=doc.title,
        manual_type=doc.manual_type,
        section=section,
        class_code=doc.class_code,
        chunk_id=_chunk_id(doc, section, ordinal, content),
        chunk_strategy=strategy,
        parser=doc.parser,
        source_hash=doc.source_hash,
        content_hash=content_hash,
        aliases=_dedupe(aliases),
        content=content,
        page=block.page,
        sheet=block.sheet,
        row_range=block.row_range,
        locator=locator,
        warnings=[*doc.warnings, *validate_chunk_shape(section, content)],
    )


def validate_chunk_shape(section: str, content: str) -> list[str]:
    warnings = []
    if len(content) < MIN_BODY_CHARS:
        warnings.append("body_too_short")
    if section and section[0].islower():
        warnings.append("section_starts_lowercase_fragment")
    if _is_toc_only(Section(section, re.sub(r"^# .+?\n+", "", content, count=1))):
        warnings.append("toc_only")
    if len(content) > DEFAULT_MAX_CHARS * 2:
        warnings.append("oversized_chunk")
    return warnings


def collect_source_files(paths: Iterable[Path | str], output_dir: Path | None = None) -> list[Path]:
    files: list[Path] = []
    output_resolved = output_dir.resolve() if output_dir else None
    for raw_path in paths:
        if _is_url(raw_path):
            continue
        path = Path(raw_path)
        if not path.exists():
            continue
        if path.is_file():
            if path.suffix.lower() in SUPPORTED_SUFFIXES:
                files.append(path)
            continue
        if path.is_dir():
            for candidate in path.rglob("*"):
                if not candidate.is_file() or candidate.suffix.lower() not in SUPPORTED_SUFFIXES:
                    continue
                if output_resolved and output_resolved in candidate.resolve().parents:
                    continue
                files.append(candidate)
    return sorted(set(files))


def collect_sources(paths: Iterable[Path | str], output_dir: Path | None = None) -> list[str | Path]:
    sources: list[str | Path] = []
    for raw_path in paths:
        if _is_url(raw_path):
            sources.append(str(raw_path))
    sources.extend(collect_source_files(paths, output_dir))
    return sources


def output_subdir(chunk: ChunkRecord) -> str:
    normalized = f"{chunk.manual_type} {chunk.source_type}".lower()
    if "property" in normalized:
        return "property"
    if "guide" in normalized:
        return "gl_guide"
    if "general liability" in normalized:
        return "gl_classes"
    if "internal" in normalized:
        return "internal"
    if "spreadsheet" in normalized:
        return "spreadsheets"
    if "pdf" in normalized:
        return "pdf"
    if "word" in normalized:
        return "word"
    if "website" in normalized:
        return "websites"
    return "other"


def write_chunks(chunks: list[ChunkRecord], output_dir: Path, *, clean: bool = False) -> list[str]:
    write_warnings: list[str] = []
    if clean and output_dir.exists():
        try:
            shutil.rmtree(output_dir, onerror=_handle_remove_readonly)
        except PermissionError as exc:
            write_warnings.append(
                "clean_failed_permission_denied: "
                f"{exc}. Existing files were overwritten where names matched."
            )
    output_dir.mkdir(parents=True, exist_ok=True)

    used_names: Counter[str] = Counter()
    for chunk in chunks:
        subdir = output_dir / output_subdir(chunk)
        subdir.mkdir(parents=True, exist_ok=True)
        base_name = f"{chunk.chunk_id}.md"
        used_names[base_name] += 1
        if used_names[base_name] > 1:
            base_name = f"{chunk.chunk_id}-{used_names[base_name]}.md"
        (subdir / base_name).write_text(chunk.document_text(), encoding="utf-8")
    return write_warnings


def _handle_remove_readonly(func, path, _exc_info) -> None:
    os.chmod(path, stat.S_IWRITE)
    func(path)


def build_manifest(
    chunks: list[ChunkRecord],
    source_refs: list[str | Path],
    documents: list[SourceDocument] | None = None,
) -> dict:
    chunk_ids_by_content_hash: dict[str, list[str]] = defaultdict(list)
    for chunk in chunks:
        chunk_ids_by_content_hash[chunk.content_hash].append(chunk.chunk_id)
    duplicate_groups = {
        content_hash: sorted(chunk_ids)
        for content_hash, chunk_ids in chunk_ids_by_content_hash.items()
        if len(chunk_ids) > 1
    }
    property_sections = sorted(
        {chunk.section for chunk in chunks if chunk.manual_type == "Property"}
    )
    chunks_by_manual_type = Counter(chunk.manual_type for chunk in chunks)
    chunks_by_source_type = Counter(chunk.source_type for chunk in chunks)
    chunks_by_strategy = Counter(chunk.chunk_strategy for chunk in chunks)
    chunks_by_section: dict[str, int] = defaultdict(int)
    for chunk in chunks:
        chunks_by_section[f"{chunk.manual_type}::{chunk.section}"] += 1

    warnings_by_chunk = {chunk.chunk_id: chunk.warnings for chunk in chunks if chunk.warnings}
    source_manifest = []
    for doc in documents or []:
        source_manifest.append(
            {
                "source_uri": doc.source_uri,
                "source_url": doc.source_url,
                "source_file": str(doc.path) if doc.path else "",
                "source_type": doc.kind,
                "manual_type": doc.manual_type,
                "title": doc.title,
                "parser": doc.parser,
                "source_hash": doc.source_hash,
                "block_count": len(doc.blocks),
                "warnings": doc.warnings,
            }
        )

    has_property_source = any(doc.manual_type == "Property" for doc in documents or []) or bool(
        property_sections
    )

    return {
        "generated_at": datetime.now(UTC).isoformat(),
        "source_count": len(source_refs),
        "source_refs": [str(ref) for ref in source_refs],
        "sources": source_manifest,
        "chunk_count": len(chunks),
        "chunks_by_manual_type": dict(sorted(chunks_by_manual_type.items())),
        "chunks_by_source_type": dict(sorted(chunks_by_source_type.items())),
        "chunks_by_strategy": dict(sorted(chunks_by_strategy.items())),
        "duplicate_content_hashes": sorted(duplicate_groups),
        "duplicate_content_groups": duplicate_groups,
        "warnings_by_chunk": warnings_by_chunk,
        "property_sections_found": property_sections,
        "missing_property_sections": (
            sorted(set(EXPECTED_PROPERTY_SECTIONS) - set(property_sections))
            if has_property_source
            else []
        ),
        "max_chunk_chars": max((len(chunk.content) for chunk in chunks), default=0),
        "chunks_by_section": dict(sorted(chunks_by_section.items())),
    }


def write_manifest(manifest: dict, output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    (output_dir / "manifest.json").write_text(
        json.dumps(manifest, indent=2, sort_keys=True),
        encoding="utf-8",
    )


def run_pipeline(
    source_paths: Iterable[Path | str],
    output_dir: Path,
    *,
    max_chars: int = DEFAULT_MAX_CHARS,
    use_agno: bool = False,
    clean: bool = False,
    strategy: str = "auto",
    prefer_agno_reader: bool = False,
) -> dict:
    source_refs = collect_sources(source_paths, output_dir)
    documents: list[SourceDocument] = []
    chunks: list[ChunkRecord] = []

    for source_ref in source_refs:
        docs = load_source_reference(source_ref, prefer_agno_reader=prefer_agno_reader)
        documents.extend(docs)
        for doc in docs:
            chunks.extend(
                build_chunks_for_document(
                    doc,
                    max_chars=max_chars,
                    strategy=strategy,
                    use_agno=use_agno,
                )
            )

    manifest = build_manifest(chunks, source_refs, documents)
    unique_chunks: list[ChunkRecord] = []
    seen_content: set[str] = set()
    for chunk in chunks:
        if chunk.content_hash in seen_content:
            continue
        seen_content.add(chunk.content_hash)
        unique_chunks.append(chunk)

    write_warnings = write_chunks(unique_chunks, output_dir, clean=clean)
    manifest["written_chunk_count"] = len(unique_chunks)
    manifest["write_warnings"] = write_warnings
    write_manifest(manifest, output_dir)
    return manifest


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description=(
            "Parse websites/files and generate local Bedrock-ready Markdown chunks. "
            "Upload output to S3 later and use Bedrock data-source chunking NONE."
        )
    )
    parser.add_argument(
        "--input",
        nargs="*",
        type=str,
        default=[str(path) for path in DEFAULT_SOURCE_PATHS],
        help="Input files, directories, or URLs.",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=DEFAULT_OUTPUT_DIR,
        help="Local output directory for generated chunk files.",
    )
    parser.add_argument("--max-chars", type=int, default=DEFAULT_MAX_CHARS)
    parser.add_argument(
        "--strategy",
        choices=["auto", "hybrid", "page_hybrid", "table", "recursive", "agentic", "semantic"],
        default="auto",
        help="Chunking strategy. auto chooses by source type.",
    )
    parser.add_argument(
        "--use-agno",
        action="store_true",
        help="Use Agno chunking for long prose when installed/configured.",
    )
    parser.add_argument(
        "--prefer-agno-reader",
        action="store_true",
        help="Try Agno ReaderFactory before deterministic parsers.",
    )
    parser.add_argument("--clean", action="store_true", help="Delete output directory first.")
    parser.add_argument(
        "--print-sources",
        action="store_true",
        help="Print source/parser summary after completion.",
    )
    return parser


def main() -> int:
    args = build_arg_parser().parse_args()
    manifest = run_pipeline(
        args.input,
        args.output,
        max_chars=args.max_chars,
        use_agno=args.use_agno,
        clean=args.clean,
        strategy=args.strategy,
        prefer_agno_reader=args.prefer_agno_reader,
    )
    print(f"Wrote {manifest['written_chunk_count']} chunks to {args.output}")
    print(f"Manifest: {args.output / 'manifest.json'}")
    if manifest["missing_property_sections"]:
        print("Missing Property sections:", ", ".join(manifest["missing_property_sections"]))
    if manifest["duplicate_content_hashes"]:
        print(f"Duplicate content hashes detected: {len(manifest['duplicate_content_hashes'])}")
    for warning in manifest.get("write_warnings", []):
        print(f"Warning: {warning}")
    if args.print_sources:
        for source in manifest.get("sources", []):
            print(
                f"- {source['source_uri']} | {source['source_type']} | "
                f"{source['manual_type']} | parser={source['parser']} | blocks={source['block_count']}"
            )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
